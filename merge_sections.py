import json
import os
from copy import deepcopy

import mammoth
from docx import Document

# ---- CONFIG ----

TEMPLATE_PATH = "master_template.docx"
SECTIONS_JSON = "sections.json"
OUTPUT_DOCX   = "merged_output.docx"

# ---- Hilfsfunktionen ----

def html_to_docx_elements(html: str):
    """
    Konvertiert ein HTML-Snippet in python-docx-Elemente mittels Mammoth.
    Liefert eine Liste von OpenXML-Elementen (paragraphs, tables, ...)
    """
    # Konvertiere HTML zu HTML-Document in-memory
    result = mammoth.convert_to_docx(html.encode("utf8"))
    docx_bytes = result.value  # das ist ein .docx in bytes
    # Lade es als python-docx Document
    with open("temp_section.docx", "wb") as f:
        f.write(docx_bytes)
    tmp = Document("temp_section.docx")
    os.remove("temp_section.docx")
    # Alle body-Elemente (paragraphs und tables) extrahieren
    return list(tmp.element.body)

# ---- Hauptlogik ----

def main():
    # 1) JSON laden
    data = json.load(open(SECTIONS_JSON, encoding="utf8"))
    sections = data["sections"]  # dict: "1"-> html, "2"-> html, …

    # 2) Template laden
    tpl = Document(TEMPLATE_PATH)
    body = tpl.element.body

    # 3) Für jede Section-Platzhalter ersetzen
    for sec, html in sections.items():
        placeholder = f"{{{{SECTION_{sec}}}}}"  # z.B. {{SECTION_1}}
        # Suche Absatz mit dem Marker
        for p in tpl.paragraphs:
            if placeholder in p.text:
                p_idx = body.index(p._p)
                parent = p._p.getparent()
                # 1) entferne den Marker-Absatz
                parent.remove(p._p)
                # 2) wandle HTML in DOCX-Elemente um
                elems = html_to_docx_elements(html)
                # 3) füge sie an gleicher Stelle ein
                for i, elm in enumerate(elems):
                    parent.insert(p_idx + i, deepcopy(elm))
                break

    # 4) Speichern
    tpl.save(OUTPUT_DOCX)
    print(f"✅ Fertig: {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()
