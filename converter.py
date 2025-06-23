import sys
import re
from typing import Dict, List
from pdf2docx import Converter
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from copy import deepcopy

# Standardpfade (werden im Container verwendet)
DEFAULT_INPUT_DIR = "/data/sample_pdfs"
DEFAULT_TEMPLATE_PATH = "/data/templates/master_template.docx"
DEFAULT_OUTPUT_DIR = "/data/output"
DEFAULT_ICONS_DIR = "/data/icons"

# Pfade ggf. überschreiben, damit das Skript auch lokal lauffähig ist
INPUT_DIR = os.environ.get("INPUT_DIR", DEFAULT_INPUT_DIR)
TEMPLATE_PATH = os.environ.get("TEMPLATE_PATH", DEFAULT_TEMPLATE_PATH)
OUTPUT_DIR = os.environ.get("OUTPUT_DIR", DEFAULT_OUTPUT_DIR)
ICONS_DIR = os.environ.get("ICONS_DIR", DEFAULT_ICONS_DIR)

print(f">>> Verwende TEMPLATE_PATH: {TEMPLATE_PATH}")

# Regex für Abschnitts-Header (flexibel für Deutsch/Englisch)
header_re = re.compile(r"^\s*(?:Abschnitt|Section)\s*\.?\s*(\d+)\s*[:\.-]?", re.I)
# Regex für Abschnitts-Header (flexibler, erlaubt Aufzählungszeichen und –)
# Beispiele: "• Abschnitt 1:", "Section 2 –", "1. Abschnitt 3- Inhalt"
header_re = re.compile(
    r"^\s*(?:[\d•*\-–.]+\s*)?(?:Abschnitt|Section)\s*\.?\s*(\d+)\s*[:.\-–]?",
    re.I,
)
# Regex für sichere Dateinamen
safe_re = re.compile(r"[^0-9A-Za-z]+")


def pdf_to_raw_docx(pdf_path: str, raw_docx_path: str) -> None:
    """
    Konvertiert PDF direkt in DOCX mit pdf2docx.
    """
    os.makedirs(os.path.dirname(raw_docx_path), exist_ok=True)
    cv = Converter(pdf_path)
    cv.convert(raw_docx_path, start=0, end=None)
    cv.close()


def iter_block_items(parent):
    """
    Generator für Absätze und Tabellen im Dokument.
    Generator für Absätze und Tabellen im gesamten Dokument.
    """
    for child in parent.element.body:
    # Durchläuft auch verschachtelte Elemente wie Textboxen
    for child in parent.element.body.iter():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def extract_sections(raw_docx_path: str) -> Dict[str, List]:
    """
    Teilt ein rohes DOCX in Abschnitte auf, markiert durch 'Abschnitt X' oder 'Section X'.
    """
    doc = Document(raw_docx_path)
    sections: Dict[str, List] = {}
    current: str = None

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            m = header_re.match(text)
            if m:
                sec = m.group(1)
                current = sec
                sections.setdefault(sec, [])
                remaining = header_re.sub("", text, count=1).strip()
                if remaining:
                    tmp_doc = Document()
                    tmp_doc.add_paragraph(remaining)
                    sections[current].append(tmp_doc.paragraphs[-1])
                continue
            if current:
                sections[current].append(block)

        elif isinstance(block, Table):
            found = False
            header_sec = None
            header_row = None
            for ri, row in enumerate(block.rows):
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if header_re.match(para.text.strip()):
                            header_sec = header_re.match(para.text.strip()).group(1)
                            found = True
                            header_row = ri
                            break
                    if found:
                        break
                if found:
                    break

            if found and header_sec:
                if header_sec != current:
                    current = header_sec
                    sections.setdefault(header_sec, [])
                tbl_elem = deepcopy(block._element)
                for _ in range(header_row + 1):
                    tbl_elem.tr_lst.pop(0)
                if tbl_elem.tr_lst:
                    sections[current].append(Table(tbl_elem, doc))
                else:
                    # Nur Kopfzeile vorhanden – komplette Tabelle übernehmen
                    sections[current].append(block)
            else:
                if current:
                    sections[current].append(block)

    return sections


def merge_into_template(sections: Dict[str, List], template_path: str, out_path: str) -> None:
    """
    Fügt die extrahierten Abschnitte in das Template an den Platzhaltern {SECTION X} ein,
    sucht dabei in allen Absätzen (inkl. Tabellenzellen).
    """
    print(f">>> merge_into_template lädt Template von: {template_path}")
    if not os.path.isfile(template_path):
        print(f"Template nicht gefunden: {template_path}", file=sys.stderr)
        return

    tpl = Document(template_path)
    # Pattern für {SECTION_1} oder {SECTION 1}
    pattern = re.compile(r"\{\s*SECTION[_ ]?(\d+)\s*\}", re.I)

    # Gehe alle w:p Knoten durch (enthält Paragraphs überall)
    for p_elem in tpl.element.body.xpath('.//w:p'):
    # Gehe alle w:p Knoten durch (auch in Textboxen und Shapes)
    for p_elem in tpl.element.xpath('.//w:p'):
        texts = [t.text for t in p_elem.xpath('.//w:t') if t.text]
        full_text = ''.join(texts)
        m = pattern.search(full_text)
        if not m:
            continue
        sec = m.group(1)
        parent = p_elem.getparent()
        idx = parent.index(p_elem)
        parent.remove(p_elem)
        for elem in sections.get(sec, []):
            new_elm = deepcopy(elem._element)
            parent.insert(idx, new_elm)
            idx += 1

    tpl.save(out_path)
    print(f">>> Template gespeichert: {out_path}")


if __name__ == '__main__':
    os.makedirs(INPUT_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(ICONS_DIR, exist_ok=True)

    for f in os.listdir(INPUT_DIR):
        if not f.lower().endswith('.pdf'):
            continue
        try:
            pdf_file = os.path.join(INPUT_DIR, f)
            base, _ = os.path.splitext(f)
            raw = os.path.join(OUTPUT_DIR, f"{base}_raw.docx")
            final = os.path.join(OUTPUT_DIR, f"{base}.docx")
            safe_base = safe_re.sub('_', base)
            raw = os.path.join(OUTPUT_DIR, f"{safe_base}_raw.docx")
            final = os.path.join(OUTPUT_DIR, f"{safe_base}.docx")

            print(f"Processing {f}...")
            pdf_to_raw_docx(pdf_file, raw)
            sections = extract_sections(raw)
            print(f">>> Debug: Sections für {f}: {list(sections.keys())}")
            for s, els in sections.items(): print(f"  Sektion {s}: {len(els)} Elemente")

            merge_into_template(sections, TEMPLATE_PATH, final)
            print(f"Saved {final}")

            if os.path.exists(pdf_file):
                os.remove(pdf_file)
                print(f"Removed input PDF: {pdf_file}")

        except Exception as exc:
            print(f"Fehler bei {f}: {exc}", file=sys.stderr)
            continue
